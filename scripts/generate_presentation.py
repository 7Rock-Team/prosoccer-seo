"""
Generate the client-facing SEO strategy presentation for Tony Tatikian.

Produces:
  - deliverables/strategy-presentations/2026-04_seo-goals-presentation-for-tony-v2.docx
  - deliverables/strategy-presentations/2026-04_seo-goals-presentation-for-tony-v2.pdf

Source content: context/06-business-goals.md (locked).

Page-layout rule: every top-level section starts on a fresh page. This is enforced
structurally by the @section decorator; don't rely on future contributors remembering
to call add_page_break() at the top of each section body.
"""

from __future__ import annotations

import functools
import os
import sys
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION

# Paths
REPO_ROOT = Path(__file__).resolve().parent.parent
LOGO_PATH = REPO_ROOT / "assets" / "brand" / "7rock-logo.png"
OUT_DIR = REPO_ROOT / "deliverables" / "strategy-presentations"
BUILD_DIR = REPO_ROOT / "scripts" / "build"
OUT_DOCX = OUT_DIR / "2026-04_seo-goals-presentation-for-tony-v2.docx"

# Brand colors
RED = "FF5151"
NEAR_BLACK = "0E0E0E"
BODY = "2C2C2C"
ACCENT_BG = "E9F6FB"
NEUTRAL_GRAY = "D9D9D9"
TEXT_GRAY = "6B6B6B"

RGB_RED = RGBColor(0xFF, 0x51, 0x51)
RGB_NEAR_BLACK = RGBColor(0x0E, 0x0E, 0x0E)
RGB_BODY = RGBColor(0x2C, 0x2C, 0x2C)
RGB_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
RGB_TEXT_GRAY = RGBColor(0x6B, 0x6B, 0x6B)

# Fonts: Inter preferred, fallback Calibri (Windows default, widely present)
HEADING_FONT = "Calibri"
BODY_FONT = "Calibri"


# =============================================================================
# CHART GENERATION
# =============================================================================

BUILD_DIR.mkdir(parents=True, exist_ok=True)
OUT_DIR.mkdir(parents=True, exist_ok=True)


def chart1_organic_share(path: Path) -> None:
    """Donut: organic share of Online Store revenue."""
    fig, ax = plt.subplots(figsize=(7, 4.2), dpi=200)
    sizes = [51, 49]
    colors = ["#FF5151", "#E9F6FB"]
    wedges, _ = ax.pie(
        sizes,
        colors=colors,
        startangle=90,
        counterclock=False,
        wedgeprops=dict(width=0.38, edgecolor="white", linewidth=2),
    )
    ax.text(0, 0.08, "51%", ha="center", va="center",
            fontsize=32, fontweight="bold", color="#0E0E0E")
    ax.text(0, -0.18, "~$950K annualized", ha="center", va="center",
            fontsize=11, color="#2C2C2C")

    ax.text(1.25, 0.6, "Google organic search", fontsize=12,
            color="#0E0E0E", fontweight="bold")
    ax.text(1.25, 0.42, "Largest single referral source", fontsize=10,
            color="#2C2C2C")
    ax.text(1.25, -0.45, "Other channels (49%)", fontsize=10,
            color="#6B6B6B")
    ax.text(1.25, -0.6, "Direct, email, paid, social, referral", fontsize=9,
            color="#6B6B6B")

    ax.set_aspect("equal")
    ax.axis("off")
    plt.tight_layout()
    fig.savefig(path, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def chart2_surface_ctr(path: Path) -> None:
    """Column chart: Merchant Listings vs Product Snippets CTR."""
    fig, ax = plt.subplots(figsize=(7, 4.2), dpi=200)
    labels = ["Merchant Listings", "Product Snippets"]
    ctrs = [6.33, 0.53]
    positions = [4.69, 21.4]
    colors = ["#FF5151", "#B8B8B8"]

    bars = ax.bar(labels, ctrs, color=colors, width=0.55, edgecolor="white")

    for bar, ctr, pos in zip(bars, ctrs, positions):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height() + 0.15,
            f"{ctr}% CTR",
            ha="center", va="bottom", fontsize=12, fontweight="bold",
            color="#0E0E0E",
        )
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height() / 2,
            f"avg pos {pos}",
            ha="center", va="center", fontsize=10, color="white",
            fontweight="bold",
        )

    ax.set_ylim(0, max(ctrs) * 1.25)
    ax.set_ylabel("Click-through rate (%)", fontsize=10, color="#2C2C2C")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color("#D9D9D9")
    ax.spines["bottom"].set_color("#D9D9D9")
    ax.tick_params(colors="#2C2C2C")
    ax.yaxis.grid(True, color="#EEEEEE", linewidth=0.8)
    ax.set_axisbelow(True)

    plt.tight_layout()
    fig.savefig(path, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def chart3_high_order_days(path: Path) -> None:
    """Column chart: high-order days (90+ orders) year over year."""
    fig, ax = plt.subplots(figsize=(7, 4.2), dpi=200)
    labels = ["2024\n(full year)", "Q3 2025 to Q1 2026\n(9 months)", "2026 target"]
    values = [30, 8, 50]
    colors = ["#0E0E0E", "#B8B8B8", "#FF5151"]

    bars = ax.bar(labels, values, color=colors, width=0.55, edgecolor="white")

    display_labels = ["30 days", "8 days", "50+ days"]
    for bar, label in zip(bars, display_labels):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height() + 1.2,
            label,
            ha="center", va="bottom", fontsize=11, fontweight="bold",
            color="#0E0E0E",
        )

    ax.set_ylim(0, 60)
    ax.set_ylabel("Days per year with 90+ orders", fontsize=10, color="#2C2C2C")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color("#D9D9D9")
    ax.spines["bottom"].set_color("#D9D9D9")
    ax.tick_params(colors="#2C2C2C")
    ax.yaxis.grid(True, color="#EEEEEE", linewidth=0.8)
    ax.set_axisbelow(True)

    plt.tight_layout()
    fig.savefig(path, bbox_inches="tight", facecolor="white")
    plt.close(fig)


# =============================================================================
# DOCX HELPERS
# =============================================================================

def set_cell_shading(cell, fill_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def remove_table_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "nil")
        borders.append(b)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


def set_cell_margins(cell, top=120, bottom=120, left=180, right=180) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for name, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{name}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def style_run(run, *, size=11, bold=False, color=RGB_BODY, font=BODY_FONT):
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:cs"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def add_paragraph(doc_or_cell, text, *, size=11, bold=False, color=RGB_BODY,
                  font=BODY_FONT, align=None, space_before=None, space_after=None,
                  line_spacing=None):
    p = doc_or_cell.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    if space_before is not None:
        pf.space_before = Pt(space_before)
    if space_after is not None:
        pf.space_after = Pt(space_after)
    if line_spacing is not None:
        pf.line_spacing = line_spacing
    run = p.add_run(text)
    style_run(run, size=size, bold=bold, color=color, font=font)
    return p


def add_heading(doc, text, *, level=1):
    # level 1 = section header, level 2 = subsection
    sizes = {0: 36, 1: 22, 2: 15, 3: 13}
    space_before = {0: 0, 1: 6, 2: 12, 3: 6}
    space_after = {0: 6, 1: 8, 2: 4, 3: 2}
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before[level])
    pf.space_after = Pt(space_after[level])
    run = p.add_run(text)
    style_run(
        run,
        size=sizes[level],
        bold=True,
        color=RGB_NEAR_BLACK,
        font=HEADING_FONT,
    )
    return p


def add_horizontal_line(paragraph, color_hex=RED, size_pt=6):
    """Add a bottom border (horizontal rule) to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size_pt * 4))  # eighths of a point
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_callout(doc, lines, *, fill=ACCENT_BG, title=None):
    """Single-cell table styled as an accent callout box.

    lines: list of either str or tuple (text, bold).
    """
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, fill)
    remove_table_borders(table)
    set_cell_margins(cell, top=200, bottom=200, left=260, right=260)

    # First paragraph already exists in the empty cell
    first = cell.paragraphs[0]
    first_used = False

    if title:
        run = first.add_run(title)
        style_run(run, size=12, bold=True, color=RGB_NEAR_BLACK, font=HEADING_FONT)
        first_used = True

    for entry in lines:
        if isinstance(entry, tuple):
            text, bold = entry
        else:
            text, bold = entry, False
        if not first_used:
            p = first
            first_used = True
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        style_run(run, size=11, bold=bold, color=RGB_NEAR_BLACK, font=BODY_FONT)

    # spacing after callout
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(0)
    return table


def add_image(doc, image_path, width_inches=6.4, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(2)
        cap.paragraph_format.space_after = Pt(12)
        crun = cap.add_run(caption)
        style_run(crun, size=9, color=RGB_TEXT_GRAY, font=BODY_FONT)
        crun.italic = True
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    from docx.enum.text import WD_BREAK
    run.add_break(WD_BREAK.PAGE)


def section(builder):
    """Decorator: enforce that a top-level section starts on a fresh page.

    Apply to every top-level section builder (Executive Summary through Next Steps).
    Do not apply to build_cover: the cover is page 1, no break before it.

    Any helper function that builds a sub-section inside a larger section should NOT
    be decorated; only the outermost builder in main() gets @section.
    """
    @functools.wraps(builder)
    def wrapped(doc, *args, **kwargs):
        add_page_break(doc)
        return builder(doc, *args, **kwargs)
    return wrapped


# Page number field helper
def insert_page_number(paragraph):
    run = paragraph.add_run()
    style_run(run, size=9, color=RGB_TEXT_GRAY, font=BODY_FONT)

    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE   \\* MERGEFORMAT "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")

    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)


def configure_footer(doc):
    """Left: 7 Rock Marketing · ProSoccer SEO Strategy 2026; Right: page number.
    Cover page (section 1) has its own empty footer.
    """
    # Use different first page footer for the whole section
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    # First page footer stays empty (cover)
    first_footer = section.first_page_footer
    first_footer.paragraphs[0].text = ""

    footer = section.footer
    p = footer.paragraphs[0]
    p.text = ""
    # Use a tab stop to right-align the page number
    from docx.enum.text import WD_TAB_ALIGNMENT
    section_width = section.page_width - section.left_margin - section.right_margin
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(section_width, WD_TAB_ALIGNMENT.RIGHT)

    left_run = p.add_run("7 Rock Marketing  ·  ProSoccer SEO Strategy 2026")
    style_run(left_run, size=9, color=RGB_TEXT_GRAY, font=BODY_FONT)
    tab_run = p.add_run("\t")
    style_run(tab_run, size=9, color=RGB_TEXT_GRAY, font=BODY_FONT)
    insert_page_number(p)


def set_page_margins(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(0.95)
        section.right_margin = Inches(0.95)


# =============================================================================
# CONTENT BUILDERS
# =============================================================================
#
# Every top-level section builder is decorated with @section, which inserts a
# page break before the section. Do not add add_page_break() calls inside a
# decorated builder's body.
# =============================================================================

def build_cover(doc):
    # Logo top-left
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_logo = p_logo.add_run()
    run_logo.add_picture(str(LOGO_PATH), width=Inches(1.6))

    # vertical spacer
    for _ in range(5):
        doc.add_paragraph()

    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_after = Pt(4)
    r_title = p_title.add_run("SEO strategy 2026")
    style_run(r_title, size=44, bold=True, color=RGB_NEAR_BLACK, font=HEADING_FONT)

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(12)
    r_sub = p_sub.add_run("Commercial plan for ProSoccer.com")
    style_run(r_sub, size=18, color=RGB_BODY, font=HEADING_FONT)

    # Red accent line
    p_line = doc.add_paragraph()
    p_line.paragraph_format.space_before = Pt(0)
    p_line.paragraph_format.space_after = Pt(18)
    add_horizontal_line(p_line, color_hex=RED, size_pt=3)

    # vertical spacer
    for _ in range(8):
        doc.add_paragraph()

    def meta_line(label, value):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run(f"{label}   ")
        style_run(r1, size=10, color=RGB_TEXT_GRAY, font=BODY_FONT)
        r2 = p.add_run(value)
        style_run(r2, size=11, bold=True, color=RGB_NEAR_BLACK, font=BODY_FONT)

    meta_line("Prepared for", "Tony Tatikian, COO, ProSoccer")
    meta_line("Prepared by", "7 Rock Marketing LLC")
    meta_line("Date", "April 2026")


@section
def build_executive_summary(doc):
    add_heading(doc, "Executive summary", level=1)

    add_paragraph(
        doc,
        "ProSoccer.com already earns most of its online store revenue from Google search. "
        "That's an asset and a risk. The commercial objective for the next twelve months is "
        "to compound the asset and cover the risk: grow organic revenue fast enough to absorb "
        "the World Cup demand spike, offset the click erosion Google's AI Overviews are starting "
        "to create, and reduce the share of online store revenue that depends on monthly Google "
        "Ads spend.",
        size=11, color=RGB_BODY,
    )

    add_callout(doc, lines=[
        ("At a glance", True),
        ("Current:  about 54 orders per day, roughly $950K/year in organic revenue, $30K/month on Google Ads.", False),
        ("12-month target:  75 to 85 orders per day, $1.2M+ in organic revenue, $5K to $8K/month in ads savings.", False),
    ])

    add_paragraph(
        doc,
        "Four goals get us there: grow non-branded organic revenue, capture World Cup 2026 "
        "traffic through an 8-week focused sprint, defend organic revenue against AI Overview "
        "click erosion through Merchant Center feed work, and establish an AI search visibility "
        "baseline while Shopify's new Agentic Storefronts feature is still young. Phase 2 "
        "discovery (April 2026) closed out the upfront diagnostic work, so Month 1 is execution, "
        "not investigation.",
        size=11, color=RGB_BODY,
    )


@section
def build_why_now(doc):
    add_heading(doc, "Why this matters now", level=1)

    add_paragraph(
        doc,
        "Roughly 51% of ProSoccer's online store revenue already comes from Google search. "
        "That's the single largest channel on the site, and it's the one with the most "
        "pull on the commercial mix. When it moves, the P&L moves with it.",
        size=11, color=RGB_BODY,
    )

    add_paragraph(
        doc,
        "Two outside forces are about to hit that channel at the same time. Google's AI "
        "Overviews now appear on about 14% of shopping queries (up from near zero six months "
        "ago) and the top organic result on those queries loses 18 to 39% of its clicks. "
        "That's the compression. Working the other direction, the FIFA World Cup opens at LA "
        "Stadium on June 12, 2026, with 78 of the 104 matches on US soil. That's the inflation. "
        "One pulls down, one pushes up, and both need a response.",
        size=11, color=RGB_BODY,
    )

    add_paragraph(
        doc,
        "One honest note on the baseline. Q1 2026 averaged roughly 54 orders per day, tracking "
        "slightly behind the pace ProSoccer held through full-year 2024 (about 57 per day). "
        "We're not relitigating that gap in this document. We're naming it so the forward plan "
        "is built on real ground, and so the targets below reflect what's actually in front of us.",
        size=11, color=RGB_BODY,
    )

    add_image(
        doc,
        BUILD_DIR / "chart1_organic_share.png",
        width_inches=6.3,
        caption=(
            "Google organic share of Online Store revenue, representative annualized figure "
            "from trailing 6 months of referrer data."
        ),
    )


@section
def build_what_we_found(doc):
    add_heading(doc, "What we found", level=1)

    add_paragraph(
        doc,
        "Phase 2 discovery ran across the first three weeks of April 2026. These are the "
        "findings that came back when we looked at the site and the Search Console data "
        "directly. They're the ground the twelve-month plan is built on.",
        size=11, color=RGB_BODY,
    )

    findings = [
        ("The ranking trajectory is improving.",
         "Average Google position has moved from 20.8 to 9.6 over the past twelve months "
         "(weekly GSC data, April 2025 to April 2026). That's not a plateau or a decline. "
         "The site is trending up."),
        ("The late-2025 theme migration caused a temporary regression, and the site has recovered.",
         "Between November 2025 and January 2026, average position slipped 3 to 4 spots, a "
         "standard post-theme-swap re-indexing effect. Rankings have since recovered to "
         "pre-migration levels across February to April 2026. Month 1 work starts from a "
         "recovered position, not a falling one. We're not claiming credit for the recovery "
         "since we haven't shipped implementation yet."),
        ("Magento legacy is clean.",
         "A scan of the 12-month GSC top-pages export for every legacy URL pattern (.html, "
         "?id=, /catalog/product, /index.php, and others) returned zero matches. No legacy "
         "URLs from the 2021-to-2022 Shopify migration are receiving traffic today. That "
         "cleanup held up. No follow-up action needed."),
        ("USMNT keyword cannibalization is a newly surfaced technical issue.",
         "Three overlapping US collection URLs are splitting link equity across the same "
         "product set. USMNT currently sits at position 46 despite a 79,559-impression pool "
         "in GSC. Canonical URL consolidation, with 301s from the non-canonical variants into "
         "one canonical, is a high-priority Month 1 technical fix and a hard prerequisite for "
         "any USMNT on-page sprint work."),
        ("National team pages aren't a single cohort. They're three different patterns.",
         "Four pages already rank in striking distance of page one (Guatemala 9.4, El Salvador "
         "10.8, South Korea 12.6, Italy 14.0). Two rank near page one despite broken or empty "
         "metadata (El Salvador 10.8, Honduras 10.7), which makes the metadata fix the "
         "lowest-effort, highest-return work in the sprint. Two pages (Mexico at 28.4, USMNT "
         "at 45.8) need heavier intervention than polish. That's why Goal 2 splits into three "
         "layers rather than treating the national team set as a uniform task."),
    ]

    for title, body in findings:
        add_heading(doc, title, level=3)
        add_paragraph(doc, body, size=11, color=RGB_BODY)


@section
def build_twelve_months_phases(doc):
    add_heading(doc, "The 12 months in three phases", level=1)

    add_paragraph(
        doc,
        "Three phases, not one linear plan. Recover what's leaking, stabilize the machine, "
        "then compound.",
        size=11, color=RGB_BODY,
    )

    phases = [
        ("Recover", "Late April to July 2026",
         "World Cup sprint execution. USMNT URL consolidation. Fast-polish and quick-win "
         "national team pages. DataFeedWatch feed audit. Agentic Storefront audit. Manual AI "
         "citation baseline. Scope is narrow and sprint-shaped. First monthly report lands end "
         "of May 2026 as the execution era's first formal deliverable."),
        ("Stabilize", "August to December 2026",
         "Post-World Cup normalization. Tier-B national team pages (confirmed by Keyword "
         "Research Agent output) enter Goal 1 rotation as compounding assets. Product schema, "
         "review schema, and Merchant Listings defense continue. Q4 holiday positioning."),
        ("Grow", "January to April 2027",
         "Compounding categories lead the workstream mix: club jerseys, goalkeeper, futsal, "
         "youth. AI search visibility tool decision revisited against 90+ days of manual "
         "tracking data. Quarterly re-plan against Q1 2027 results."),
    ]

    table = doc.add_table(rows=1, cols=3)
    table.autofit = True
    remove_table_borders(table)

    for i, (name, months, desc) in enumerate(phases):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, ACCENT_BG)
        set_cell_margins(cell, top=240, bottom=240, left=240, right=240)

        # Phase name in red
        cell.paragraphs[0].text = ""
        p_name = cell.paragraphs[0]
        r_name = p_name.add_run(name)
        style_run(r_name, size=16, bold=True, color=RGB_RED, font=HEADING_FONT)

        p_months = cell.add_paragraph()
        p_months.paragraph_format.space_after = Pt(8)
        r_m = p_months.add_run(months)
        style_run(r_m, size=10, bold=True, color=RGB_NEAR_BLACK, font=BODY_FONT)

        p_desc = cell.add_paragraph()
        r_d = p_desc.add_run(desc)
        style_run(r_d, size=10, color=RGB_BODY, font=BODY_FONT)


@section
def build_four_goals(doc):
    add_heading(doc, "The four primary goals", level=1)

    # Goal 1
    add_heading(doc, "1. Grow non-branded organic revenue", level=2)
    add_paragraph(
        doc,
        "Push the queries ProSoccer already ranks for (positions 11 to 30) onto page one.",
        size=11, bold=True, color=RGB_NEAR_BLACK,
    )
    add_paragraph(
        doc,
        "About 3.7M non-branded impressions are currently sitting at positions 11 through 30 "
        "in Google. These are queries ProSoccer already ranks for; they just don't click yet "
        "because they're on page two or three. Moving a meaningful share of them into the top "
        "ten is the single largest organic lift available, and it's the lift that most directly "
        "reduces Google Ads dependency (new customers, not people searching for our brand).",
        size=11, color=RGB_BODY,
    )
    add_paragraph(
        doc,
        "Target categories to validate in Month 1: club team jerseys (Real Madrid, Barcelona, "
        "Manchester City, Liverpool, Juventus, Bayern, PSG, plus LA-relevant MLS clubs), "
        "goalkeeper gloves and keeper cleats, futsal and indoor shoes, youth cleats and sizing "
        "content, national team jerseys with LA-area diaspora demand, and position-specific "
        "cleat content. Final priority order comes out of a Category Priority Matrix built from "
        "Shopify sales data, GSC queries, and seasonality.",
        size=11, color=RGB_BODY,
    )
    add_callout(doc, lines=[
        ("Why lead with this: branded queries already rank 1 to 3 and are stable. Growth has "
         "to come from non-branded, and non-branded is where ads dependency actually breaks.", False),
    ])

    # Goal 2
    add_heading(doc, "2. Capture World Cup 2026 traffic", level=2)
    add_paragraph(
        doc,
        "An 8-week sprint tied to the tournament, narrow scope, close to the catalog.",
        size=11, bold=True, color=RGB_NEAR_BLACK,
    )
    add_paragraph(
        doc,
        "The FIFA World Cup runs June 11 through July 19, 2026. The opener is at LA Stadium "
        "on June 12, inside ProSoccer's home market. Phase 2 discovery showed the national "
        "team collection pages fall into three layers, and each layer has a different workload "
        "inside the sprint.",
        size=11, color=RGB_BODY,
    )

    add_heading(doc, "Layer 1, fast polish (4 pages)", level=3)
    add_paragraph(
        doc,
        "Guatemala (position 9.4), Italy (14.0), South Korea (12.6), and El Salvador (10.8). "
        "Standard on-page optimization: titles, meta descriptions, H1 polish, and intro copy "
        "that earns the click. El Salvador's broken-metadata fix folds into this page's pass, "
        "not a separate workstream. These four represent the highest-probability page-one gains "
        "available inside the 8-week window.",
        size=11, color=RGB_BODY,
    )

    add_heading(doc, "Layer 2, heavy lift (2 pages)", level=3)
    add_paragraph(
        doc,
        "Mexico (currently at position 28.4) needs an on-page rewrite plus internal linking, "
        "not polish. USMNT (45.8) requires canonical URL consolidation before any on-page work "
        "will pay off. The consolidation is a Technical SEO task and a hard prerequisite for "
        "the USMNT content sprint. Mexico carries the bigger opportunity on paper (119,131 GSC "
        "impressions versus USMNT's 79,559) and the LA Mexican diaspora strength; expect "
        "measurable progress inside the sprint, not full recovery.",
        size=11, color=RGB_BODY,
    )

    add_heading(doc, "Layer 3, quick win (1 page)", level=3)
    add_paragraph(
        doc,
        "Honduras (position 10.7) with empty metadata and only 6 products still pulls a 1.15% "
        "click-through rate from a small impression pool. Writing a proper title and meta "
        "description is the entire scope. Lowest-effort, highest-return item on the sprint board.",
        size=11, color=RGB_BODY,
    )

    add_paragraph(
        doc,
        "Plus two non-country pieces: one local \"where to watch in LA\" guide tied to the "
        "Pasadena store, and one authenticity guide (real vs. replica vs. fake jersey) that "
        "converts casual fans.",
        size=11, color=RGB_BODY,
    )
    add_paragraph(
        doc,
        "One pending follow-up may reshuffle Layer 1 priority: which countries actually play "
        "matches at LA Stadium during the tournament isn't yet confirmed. Once FIFA publishes "
        "the venue-by-venue schedule, we'll re-check Layer 1 against the LA-hosted countries "
        "and may promote or swap pages based on which national teams are competing in "
        "ProSoccer's home market.",
        size=11, color=RGB_BODY,
    )
    add_paragraph(
        doc,
        "What we won't chase: generic \"World Cup 2026\" queries. FIFA.com, Fox Sports, and ESPN "
        "take that traffic; we'd lose. Narrow and close to the catalog wins.",
        size=11, color=RGB_BODY,
    )
    add_callout(doc, lines=[
        ("Urgency: 8 weeks to opening match. Any approval slippage compresses the window "
         "further, so we'd like a weekly batch approval model for this sprint.", False),
    ])

    # Goal 3
    add_heading(doc, "3. Defend organic revenue against AI Overview click erosion", level=2)
    add_paragraph(
        doc,
        "Get more of the catalog into Google's Merchant Listings surface, cleanly.",
        size=11, bold=True, color=RGB_NEAR_BLACK,
    )
    add_paragraph(
        doc,
        "Not all Google search surfaces are equally affected by AI Overviews. In ProSoccer's "
        "own data, Google's Merchant Listings surface pulls a 6.33% click-through rate at an "
        "average position of 4.69. Google's Product Snippets surface pulls 0.53% at position "
        "21.4. Merchant Listings are roughly 12 times more click-efficient per impression, and "
        "they sit in a spot that AI Overviews haven't cannibalized.",
        size=11, color=RGB_BODY,
    )
    add_paragraph(
        doc,
        "The named workstream: Merchant Center feed optimization via DataFeedWatch. Every "
        "product title, attribute, image, GTIN, availability signal, and review integration "
        "in that feed is a lever on Merchant Listings performance. 7 Rock already manages "
        "the DataFeedWatch instance, which means the audit and fixes start immediately. This "
        "is likely the fastest-moving lever on the whole engagement: feed improvements show "
        "up in Merchant Listings within weeks, not quarters.",
        size=11, color=RGB_BODY,
    )
    add_callout(doc, lines=[
        ("Why defensive, not offensive: Goal 1 grows new traffic. Goal 3 protects the "
         "traffic we already have. Letting a Q3 report read \"impressions surged, revenue "
         "flat\" isn't an acceptable outcome.", False),
    ])

    add_image(
        doc,
        BUILD_DIR / "chart2_surface_ctr.png",
        width_inches=6.3,
        caption=(
            "Google surface performance: where ProSoccer products actually click. "
            "Merchant Listings convert 12x more efficiently than organic product pages."
        ),
    )

    # Goal 4
    add_heading(doc, "4. Establish AI search visibility baseline", level=2)
    add_paragraph(
        doc,
        "Measure what's already happening and position the catalog for the shift that's coming.",
        size=11, bold=True, color=RGB_NEAR_BLACK,
    )
    add_paragraph(
        doc,
        "AI search engines (ChatGPT, Perplexity, Gemini, Copilot) now handle an estimated 12 "
        "to 18% of English informational queries. Shopify reported a 15x increase in "
        "AI-originated orders across its merchant base over the past 12 months. ChatGPT is "
        "already sending real orders to ProSoccer.com: 9 orders for about $1,006 in Q1 2026, "
        "10 orders for $1,367 in Q4 2025. Small numbers, still real, still trending with the "
        "broader shift.",
        size=11, color=RGB_BODY,
    )
    add_paragraph(
        doc,
        "In late March 2026, Shopify activated Agentic Storefronts for all stores as part of "
        "the Winter '26 Edition. ProSoccer's products are auto-syndicated to Google AI Mode, "
        "ChatGPT, Microsoft Copilot, and Perplexity through the Shopify Catalog by default. "
        "Month 1 work: audit the ProSoccer Shopify admin to confirm which AI channels are "
        "enabled, verify store policies and product attributes are clean (AI platforms match "
        "on structured data, not images), and build a manual citation tracking sheet against "
        "~20 target prompts across ChatGPT, Perplexity, and Gemini.",
        size=11, color=RGB_BODY,
    )
    add_callout(doc, lines=[
        ("No AI-visibility tool purchase recommended yet. 90 days of manual tracking "
         "establishes which prompts and categories actually matter. Tool decision revisits at "
         "end of Q3 2026.", False),
    ])


@section
def build_execution_start(doc):
    add_heading(doc, "Late April to May 2026: execution start", level=1)

    add_paragraph(
        doc,
        "Phase 2 discovery (April 2026) answered most of the diagnostic questions we would "
        "otherwise have opened Month 1 with. We don't need an investigation to tell us whether "
        "the traffic problem is real. We need execution to act on what we already know. Here's "
        "what the data says and how that shapes the first weeks of work.",
        size=11, color=RGB_BODY,
    )

    findings = [
        ("Traffic trajectory: improving.",
         "Average Google position moved from 20.8 to 9.6 over the past twelve months. Clicks "
         "grew roughly 58% in the same window. Traffic is not the problem."),
        ("Technical integrity: clean on the legacy-URL front.",
         "No Magento-era URLs are receiving impressions or clicks in current GSC data. The "
         "2021-to-2022 migration cleanup held up."),
        ("Theme migration impact: temporary, and resolved.",
         "The late-2025 theme ship caused a 3-to-4 position regression across November 2025 "
         "to January 2026. Rankings have since recovered to pre-migration levels. Any residual "
         "theme-level SEO issues are now affecting the recovered baseline, which means fixes "
         "produce lift above current position, not just a return to neutral."),
        ("Conversion gap hypothesis: still valid.",
         "Q1 2026 averaged roughly 54 orders per day, below what the current impression "
         "trajectory would predict. Q1 2026 Online Store orders dropped 22% versus Q4 2025 "
         "even as average position improved. The 5.5-position gap between desktop (19.4) and "
         "mobile (13.9) is a plausible contributor, possibly a rendering or template issue "
         "specific to desktop. Continued conversion work flows into the separate "
         "prosoccer-theme CRO project that 7 Rock already manages, not this SEO scope. The "
         "SEO team flags conversion issues when the data reveals them; it doesn't own the fixes."),
        ("New technical finding: USMNT cannibalization.",
         "Three overlapping US collection URLs are splitting link equity. Canonical "
         "consolidation is named as a high-priority Month 1 deliverable and a hard "
         "prerequisite for the USMNT on-page sprint work."),
    ]

    for title, body in findings:
        add_heading(doc, title, level=3)
        add_paragraph(doc, body, size=11, color=RGB_BODY)

    add_callout(doc, lines=[
        ("Month 1 scope shifts accordingly: less diagnose, more ship. The Technical SEO "
         "Month 1 list now explicitly includes USMNT URL consolidation as a named "
         "high-priority deliverable. Device-level conversion findings continue to flow to "
         "the CRO project via Mike's mike-audit branch in the theme repo.", False),
    ])


@section
def build_outcomes(doc):
    add_heading(doc, "The outcomes we'll measure", level=1)

    add_paragraph(
        doc,
        "Every monthly report leads with these four metrics. Secondary metrics follow for "
        "program-health context.",
        size=11, color=RGB_BODY,
    )

    # Primary outcomes table
    headers = ["Metric", "Current baseline", "12-month target", "Source"]
    rows = [
        ["Online orders per day (in-scope channels)", "~54", "75 to 85", "Shopify (excl. POS, Channable, Draft Orders)"],
        ["Google organic revenue (rolling 12 months)", "~$950K", "$1.2M+", "Shopify search/google referrer"],
        ["Google Ads monthly spend reduction", "$30K/month baseline", "$5K to $8K/month saved (17-27%)", "Tony's Ads account"],
        ["High-order days per year (90+ orders)", "~8", "50+", "Shopify, in-scope channels"],
    ]

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.autofit = True
    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr_cells[i], ACCENT_BG)
        set_cell_margins(hdr_cells[i], top=140, bottom=140, left=160, right=160)
        hdr_cells[i].paragraphs[0].text = ""
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(h)
        style_run(r, size=10, bold=True, color=RGB_NEAR_BLACK, font=HEADING_FONT)

    # Body rows
    for ri, row in enumerate(rows, start=1):
        cells = table.rows[ri].cells
        for ci, val in enumerate(row):
            set_cell_margins(cells[ci], top=120, bottom=120, left=160, right=160)
            cells[ci].paragraphs[0].text = ""
            p = cells[ci].paragraphs[0]
            r = p.add_run(val)
            bold = ci in (1, 2)  # emphasize current/target columns
            style_run(r, size=10, bold=bold, color=RGB_BODY, font=BODY_FONT)

    # Light bottom border row-separator
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "D9D9D9")
        borders.append(b)
    for edge in ("insideH",):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "EEEEEE")
        borders.append(b)
    inside_v = OxmlElement("w:insideV")
    inside_v.set(qn("w:val"), "nil")
    borders.append(inside_v)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)

    doc.add_paragraph()

    add_paragraph(
        doc,
        "About the 75 to 85 per day range: 75 is the defensible floor we expect to hit, a 39% "
        "lift on the current pace delivered by the 3.7M non-branded impression pool, the World "
        "Cup sprint, and Merchant Listings defense. 85 is the stretch target; it annualizes to "
        "roughly 31,000 in-scope orders, which would clear ProSoccer's 2024 full-year total. "
        "Compounding organic channels scale non-linearly once the foundation is right, so 85 "
        "is credible as ambition, not as commitment. Any conversion-rate fixes shipped through "
        "the separate prosoccer-theme CRO project add upside on top of that range.",
        size=10, color=RGB_BODY,
    )

    add_image(
        doc,
        BUILD_DIR / "chart3_high_order_days.png",
        width_inches=6.3,
        caption=(
            "2024 = 30 high-order days (per Shopify channel data). Current 9-month window "
            "tracking at 8. 2026 target of 50+ days rebuilds on the 2024 baseline."
        ),
    )

    add_heading(doc, "Secondary metrics tracked monthly", level=3)
    sec_bullets = [
        "Keyword rankings, priority set of 50 to 100 commercial keywords, tracked weekly.",
        "Organic sessions and organic conversion rate, GA4, month over month.",
        "Indexed pages and technical health, Screaming Frog crawl, quarterly.",
        "Core Web Vitals, mobile and desktop, tracked monthly via PageSpeed Insights.",
        "Backlink profile changes via Ahrefs (pending tool approval).",
        "Merchant Listings click share versus organic click share, Search Console, monthly.",
        "AI search citation count from the manual tracking sheet in Goal 4.",
    ]
    for b in sec_bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(b)
        style_run(run, size=10, color=RGB_BODY, font=BODY_FONT)


@section
def build_what_we_wont_do(doc):
    add_heading(doc, "What we won't do", level=1)

    add_paragraph(
        doc,
        "Keep this list short. Reference it whenever scope creep shows up.",
        size=11, color=RGB_BODY,
    )

    items = [
        "No generic \"sports store\" or \"athletic gear\" optimization. ProSoccer is soccer-specific; pretending otherwise dilutes everything.",
        "No bottom-of-funnel discount queries that erode margin (\"soccer cleats under $30\" and similar).",
        "No local SEO in markets where ProSoccer has no storefront and no service coverage.",
        "No generic World Cup content where we'd lose to FIFA.com, Fox Sports, and ESPN. Specific national team pages, specific LA viewing content, and authenticity content inside the catalog context are all in scope.",
        "No content published without a clear tie to a goal on this page.",
        "No replacing Google Ads. The plan reduces its share of the revenue mix; it doesn't kill the program.",
    ]
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        style_run(run, size=11, color=RGB_BODY, font=BODY_FONT)


@section
def build_starting_assumptions(doc):
    add_heading(doc, "Starting assumptions", level=1)

    add_paragraph(
        doc,
        "These are the inputs anchoring every target and metric in this document. If any have "
        "shifted since we last aligned, flag them so we can recalibrate before Month 1 ships.",
        size=11, color=RGB_BODY,
    )

    confirmed = [
        ("Google Ads monthly spend baseline", "$30,000 per month, all campaigns."),
        ("Order volume scope",
         "All sales channels except Point of Sale, Channable (Amazon/eBay), and Draft Orders. "
         "Online Store, Shop, Tapcart, BSS B2B, Redo, and other web channels are all counted."),
        ("High-order day threshold", "90+ orders per day, using the scope above."),
        ("Affiliate program operation",
         "7 Rock manages the ProSoccer affiliate program on AWIN (Affiliate Window) and uses "
         "PayAudit as the pre-AWIN transaction approval system. Any backlink audit or disavow "
         "recommendation cross-references both rosters before action."),
    ]

    for label, value in confirmed:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r_label = p.add_run(f"{label}.  ")
        style_run(r_label, size=11, bold=True, color=RGB_NEAR_BLACK, font=BODY_FONT)
        r_val = p.add_run(value)
        style_run(r_val, size=11, color=RGB_BODY, font=BODY_FONT)


@section
def build_open_questions(doc):
    add_heading(doc, "Open questions for your review", level=1)

    add_paragraph(
        doc,
        "Three items need your input before Month 1 locks. Everything else on the plan is "
        "proceeding on the stated assumptions.",
        size=11, color=RGB_BODY,
    )

    def render_q(num, title, body, *, space_after=8):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        r_num = p.add_run(f"{num}.  ")
        style_run(r_num, size=11, bold=True, color=RGB_RED, font=BODY_FONT)
        r_title = p.add_run(title + "  ")
        style_run(r_title, size=11, bold=True, color=RGB_NEAR_BLACK, font=BODY_FONT)
        r_body = p.add_run(body)
        style_run(r_body, size=11, color=RGB_BODY, font=BODY_FONT)

    render_q(1, "Attribution model.",
        "Our baselines use Shopify's \"search/google\" referrer as the organic signal. Do you "
        "track organic revenue under a specific attribution model (last-click, assisted, "
        "post-view) we should align with, or should we additionally pull GA4 Organic Search "
        "data to triangulate?")

    # Question 2 has a multi-part body: lead-in, two bullets, trailing paragraph.
    render_q(2, "World Cup sprint approval model.",
        "The 8-week sprint covers 9 pieces of work total:",
        space_after=4)

    q2_bullets = [
        "Optimization of 7 existing national team collection pages (we're not creating new "
        "collections; these already exist on ProSoccer). The work spans 3 layers as described "
        "in Goal 2: fast polish on pages already ranking in striking distance, heavier work on "
        "pages that need rebuilding, and quick metadata fixes on pages that rank despite broken "
        "titles and descriptions.",
        "Two new blog articles: an LA watch guide and an authenticity guide.",
    ]
    for b in q2_bullets:
        bp = doc.add_paragraph(style="List Bullet")
        bp.paragraph_format.space_after = Pt(4)
        bp.paragraph_format.left_indent = Inches(0.6)
        run = bp.add_run(b)
        style_run(run, size=11, color=RGB_BODY, font=BODY_FONT)

    p_tail = doc.add_paragraph()
    p_tail.paragraph_format.space_after = Pt(8)
    p_tail.paragraph_format.left_indent = Inches(0.3)
    r_tail = p_tail.add_run(
        "Given the timeline, we'd like to run weekly batch approvals rather than "
        "page-by-page sign-off. Is that okay with you?"
    )
    style_run(r_tail, size=11, color=RGB_BODY, font=BODY_FONT)

    render_q(3, "Cadence confirmation.",
        "Monthly reports starting end of May 2026, delivered from reports/monthly/. Format "
        "and medium preferences: PDF, Google Doc, live walkthrough?")


@section
def build_next_steps(doc):
    add_heading(doc, "Next steps", level=1)

    steps = [
        "Review this document and note any changes.",
        "30-minute alignment call to answer the three open questions.",
        "Upon approval, Month 1 execution begins: USMNT URL consolidation, the fast-polish "
        "layer of the World Cup sprint, and the DataFeedWatch feed audit.",
        "First monthly report delivered end of May 2026.",
    ]
    for i, s in enumerate(steps, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r_num = p.add_run(f"{i}.  ")
        style_run(r_num, size=12, bold=True, color=RGB_RED, font=BODY_FONT)
        r_text = p.add_run(s)
        style_run(r_text, size=11, color=RGB_BODY, font=BODY_FONT)

    doc.add_paragraph()
    p_close = doc.add_paragraph()
    p_close.paragraph_format.space_before = Pt(16)
    r = p_close.add_run("Prepared to move on your signal.")
    style_run(r, size=12, bold=True, color=RGB_NEAR_BLACK, font=HEADING_FONT)

    # Red accent line below closing
    p_line = doc.add_paragraph()
    p_line.paragraph_format.space_before = Pt(4)
    add_horizontal_line(p_line, color_hex=RED, size_pt=2)


# =============================================================================
# MAIN
# =============================================================================

def main():
    print("Generating charts...")
    chart1_organic_share(BUILD_DIR / "chart1_organic_share.png")
    chart2_surface_ctr(BUILD_DIR / "chart2_surface_ctr.png")
    chart3_high_order_days(BUILD_DIR / "chart3_high_order_days.png")

    print("Building document...")
    doc = Document()

    # Base style: Calibri
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(11)

    set_page_margins(doc)

    # Build pages in order. Every @section-decorated builder inserts its own
    # page break before rendering, so section ordering here IS the page order.
    build_cover(doc)
    build_executive_summary(doc)
    build_why_now(doc)
    build_what_we_found(doc)
    build_twelve_months_phases(doc)
    build_four_goals(doc)
    build_execution_start(doc)
    build_outcomes(doc)
    build_what_we_wont_do(doc)
    build_starting_assumptions(doc)
    build_open_questions(doc)
    build_next_steps(doc)

    # Configure footer last (after sections are set)
    configure_footer(doc)

    print(f"Saving {OUT_DOCX}")
    doc.save(OUT_DOCX)
    print("Done.")


if __name__ == "__main__":
    main()
