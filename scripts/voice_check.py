"""Voice-check any file against the forbidden list in context/03-brand-voice.md.

Usage:
    python scripts/voice_check.py <path>
    python scripts/voice_check.py --file <path>

Supports .md, .txt, and .docx. Exit code 0 on clean, non-zero on violations.
"""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

FORBIDDEN_WORDS = [
    "delve", "delving",
    "unlock", "unlocks", "unlocked",
    "elevate", "elevating",
    "revolutionize",
    "seamless",
    "cutting-edge",
    "game-changer",
    "unleash",
]

FORBIDDEN_PHRASES = [
    "leverage",   # as verb; enforced with \b boundaries below
    "leverages",
    "leveraging",
    "in today's world",
    "it's important to note",
    "navigate the complex landscape",
    "dive into",
    "embark on a journey",
]

SENTENCE_OPENERS = ["In conclusion", "In summary"]

EM_DASHES = ["—", "–"]  # — and –

# Brand styling: "adidas" is always lowercase, even at sentence start (official
# trademark styling). Flag a capitalized, standalone "Adidas", EXCEPT when it is:
#   - inside a fenced block or inline backticks (already blanked by strip_backticks)
#   - immediately preceded by a backtick or a hyphen (taxonomy compounds like
#     "non-Adidas", and any backtick-adjacent literal that survived stripping)
#   - immediately followed by a hyphen + word char ("Adidas-only", "Adidas-licensed",
#     "Adidas-specific" -- workforce taxonomy jargon, not output styling)
#   - on a line carrying a pedagogical anti-pattern marker (see PEDAGOGICAL_MARKERS)
ADIDAS_PATTERN = re.compile(r"(?<![-`])\bAdidas\b(?!-\w)")

# US market language: soccer footwear is "cleats" (US market) not "boots"
# (UK/global convention). Flag standalone "boot"/"boots" in body copy; the US
# market term is "cleat(s)" (primary), "shoe(s)" (acceptable secondary variation).
# Non-soccer technical and idiom uses ("boot up", "boot camp", "boot loader",
# "to boot", "das boot") are blanked before the check so they do not false-flag;
# "bootstrap" and "reboot" never match because \b requires a word boundary.
# Backtick/fenced content and pedagogical anti-pattern lines are exempt, same as
# the Adidas check above.
BOOT_PATTERN = re.compile(r"\bboots?\b", re.IGNORECASE)
BOOT_NON_SOCCER = re.compile(
    r"\bboots?\s+(?:up|camp|loader|sector|disk|drive|menu|partition|record|screen)\b"
    r"|\bto\s+boot\b"
    r"|\bdas\s+boot\b"
    r"|\bboots?\s+on\s+the\s+ground\b",
    re.IGNORECASE,
)

# Internal link format: link suggestions in deliverables and briefings must be full
# HTTPS URLs on the canonical domain (https://www.prosoccer.com/...). These two
# patterns catch the two failure modes seen in production: an insecure http://
# ProSoccer URL, and a mangled link that lost its domain segment (http:///path).
# These checks run ONLY when the file path is inside deliverables/ or briefings/;
# the playbooks carry these strings as pedagogical INCORRECT examples and are out of
# scope. Backtick/fenced content and pedagogical lines are exempt, same as above.
LINK_INSECURE_PATTERN = re.compile(r"http://[^/\s]*prosoccer\.com", re.IGNORECASE)
LINK_MANGLED_PATTERN = re.compile(r"http:/{3,}[a-z]", re.IGNORECASE)

# Lines demonstrating what NOT to do are exempt from the Adidas check; the
# capitalized form on these lines is intentional teaching content.
PEDAGOGICAL_MARKERS = (
    "incorrect",
    "do not use",
    "anti-pattern",
    "stuffed:",
    "wrong:",
    "bad example",
    "forbidden",
    "uk convention",
    "uk-convention",
    "uk/global",
)

# Editorial body H2 casing (codified 2026-06-17, e6bdec9; voice-check enforcement
# added 2026-06-29 after Batch 4 / KK3725 surfaced the gap). Editorial body H2s --
# the H2 sections between the "### Description" marker and the "## Product Details"
# marker -- use SENTENCE case: the first word is capitalized. This is a
# scope-limited check: it flags only LOWERCASE-INITIAL body H2s. It does NOT
# attempt reverse Title-Case-drift detection (that stays at SCRIBE Phase 4 + ORIN
# Gate 15 per honest scope discipline). Because the sole permitted lowercase
# opener is the explicit "adidas" brand exception, the check cannot false-positive
# on other brand tokens (F50, Nike, FG, Gripknit) -- those are capitalized or
# uppercase and pass. Only PDP/collection briefs carry both region markers, so
# playbooks and briefings never form a region and never match. Canonical rule:
# context/page-type-playbooks/product-page-playbook.md 'H2 title casing: split discipline'.
#
# Heading-level-agnostic (added 2026-06-30 after the KI0586 Copa Elite exemplar
# surfaced a blind spot): the markers and the body-header pattern match heading
# levels 2 through 5 (## .. #####), not just ##. KI0586 used #### body sections
# (one level too deep) with lowercase first words and slipped past the original
# ##-only check. The Description -> Product Details region and the editorial body
# headers are now detected regardless of heading depth. Regression: see
# scripts/test_voice_check.py test_5_ki0586_h4_level_headers_regression.
BODY_H2_REGION_START = re.compile(r"^#{2,4}\s+Description\b", re.IGNORECASE)
BODY_H2_REGION_END = re.compile(r"^#{2,5}\s+Product\s+Details\b", re.IGNORECASE)
BODY_H2_PATTERN = re.compile(r"^#{2,5}\s+(\S+)")
ADIDAS_H2_EXCEPTION = "adidas"

SUPPORTED_EXTS = {".md", ".txt", ".docx"}


def extract_text(path: Path) -> str:
    """Return the full readable text of the file, format-aware."""
    suffix = path.suffix.lower()
    if suffix in {".md", ".txt"}:
        return path.read_text(encoding="utf-8")
    if suffix == ".docx":
        try:
            from docx import Document
        except ImportError:
            print("ERROR: python-docx is not installed. Run: pip install python-docx", file=sys.stderr)
            sys.exit(2)
        doc = Document(path)
        parts: list[str] = []
        for p in doc.paragraphs:
            parts.append(p.text)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        parts.append(p.text)
        for section in doc.sections:
            for hdr_ftr in (section.header, section.footer, section.first_page_header, section.first_page_footer):
                for p in hdr_ftr.paragraphs:
                    parts.append(p.text)
        return "\n".join(parts)
    raise ValueError(f"Unsupported file type: {suffix}. Supported: {sorted(SUPPORTED_EXTS)}")


def strip_backticks(text: str) -> str:
    """Blank out content inside fenced code blocks (```...```) and inline backticks (`...`).

    Voice rules apply to ProSoccer's own prose, not to verbatim quotations of
    competitor copy, code, or terminal output that get cited inside backticks.
    Replacement preserves newlines so line numbers in violation output stay accurate.
    """
    def blank_keep_newlines(match: re.Match) -> str:
        return "".join("\n" if c == "\n" else " " for c in match.group(0))

    text = re.sub(r"```.*?```", blank_keep_newlines, text, flags=re.DOTALL)
    text = re.sub(r"`[^`\n]*`", blank_keep_newlines, text)
    return text


def find_line_contexts(
    scan_lines: list[str], display_lines: list[str], pattern: re.Pattern
) -> list[tuple[int, str]]:
    """Return (line_number, line_text) tuples for each line where the pattern matches.

    Matches are evaluated against scan_lines (backtick content blanked); the
    displayed context is taken from display_lines (original text).
    """
    hits: list[tuple[int, str]] = []
    for idx, (sline, dline) in enumerate(zip(scan_lines, display_lines), start=1):
        if pattern.search(sline):
            hits.append((idx, dline.strip()))
    return hits


def find_lowercase_body_h2s(lines: list[str]) -> list[tuple[int, str]]:
    """Return (line_number, line_text) for editorial body H2s whose first word is
    lowercase, the "adidas" brand exception aside.

    The editorial body region is bounded by the "### Description" marker (start)
    and the first following "## Product Details" marker (end). Only product /
    collection briefs carry both markers, so other files never form a region and
    never produce a hit. Pedagogical anti-pattern lines are skipped, consistent
    with the other checks. Structural H2s ("Product Details:", "Care and
    Maintenance", "FAQs about ...") and FAQ H3 questions sit outside the region or
    are already Title/sentence case starting uppercase, so they never match here.
    """
    hits: list[tuple[int, str]] = []
    in_region = False
    for idx, line in enumerate(lines, start=1):
        if not in_region:
            if BODY_H2_REGION_START.search(line):
                in_region = True
            continue
        if BODY_H2_REGION_END.search(line):
            break
        lowered = line.lower()
        if any(marker in lowered for marker in PEDAGOGICAL_MARKERS):
            continue
        m = BODY_H2_PATTERN.match(line)
        if not m:
            continue
        first_word = m.group(1)
        if first_word == ADIDAS_H2_EXCEPTION:
            continue
        first_char = first_word[0]
        if first_char.isalpha() and first_char.islower():
            hits.append((idx, line.strip()))
    return hits


def collect_violations(text: str, path: Path | None = None) -> list[str]:
    """Return the list of voice-check violation strings (empty list if clean).

    This is the reusable core of the voice check: it performs every detection and
    returns the human-readable violation lines without printing or exiting. `check()`
    wraps it for CLI use; `scripts/batch_gate.py` imports it to fold the voice checks
    (em-dash, forbidden words/phrases/openers, capitalized Adidas, UK boots, lowercase
    editorial body H2s, internal-link URL format) into the batch gate report without
    duplicating any of that logic.

    When `path` is inside deliverables/ or briefings/, also enforce the internal-link
    URL-format checks (full HTTPS canonical-domain links; no insecure or mangled URLs).
    """
    violations: list[str] = []
    display_lines = text.splitlines()
    scan_lines = strip_backticks(text).splitlines()
    link_scope = path is not None and bool(
        re.search(r"deliverables|briefings", str(path), re.IGNORECASE)
    )

    for em in EM_DASHES:
        pattern = re.compile(re.escape(em))
        hits = find_line_contexts(scan_lines, display_lines, pattern)
        if hits:
            name = "em-dash" if em == "—" else "en-dash"
            violations.append(f"{name.upper()} '{em}' found on {len(hits)} line(s):")
            for ln, ctx in hits[:5]:
                violations.append(f"    line {ln}: {ctx[:120]}")
            if len(hits) > 5:
                violations.append(f"    ... and {len(hits) - 5} more")

    for word in FORBIDDEN_WORDS:
        pattern = re.compile(r"\b" + re.escape(word) + r"\b", re.IGNORECASE)
        hits = find_line_contexts(scan_lines, display_lines, pattern)
        if hits:
            violations.append(f"FORBIDDEN WORD '{word}' found on {len(hits)} line(s):")
            for ln, ctx in hits[:5]:
                violations.append(f"    line {ln}: {ctx[:120]}")
            if len(hits) > 5:
                violations.append(f"    ... and {len(hits) - 5} more")

    for phrase in FORBIDDEN_PHRASES:
        pattern = re.compile(r"\b" + re.escape(phrase) + r"\b", re.IGNORECASE)
        hits = find_line_contexts(scan_lines, display_lines, pattern)
        if hits:
            violations.append(f"FORBIDDEN PHRASE '{phrase}' found on {len(hits)} line(s):")
            for ln, ctx in hits[:5]:
                violations.append(f"    line {ln}: {ctx[:120]}")
            if len(hits) > 5:
                violations.append(f"    ... and {len(hits) - 5} more")

    for opener in SENTENCE_OPENERS:
        pattern = re.compile(r"(?:^|[.!?]\s+)" + re.escape(opener), re.MULTILINE)
        hits = find_line_contexts(scan_lines, display_lines, pattern)
        if hits:
            violations.append(f"FORBIDDEN SENTENCE OPENER '{opener}' found on {len(hits)} line(s):")
            for ln, ctx in hits[:5]:
                violations.append(f"    line {ln}: {ctx[:120]}")

    # Brand styling check: capitalized "Adidas" (adidas is always lowercase).
    # Skip lines carrying a pedagogical anti-pattern marker; backtick/fenced
    # content and taxonomy compounds are excluded by ADIDAS_PATTERN itself.
    adidas_hits: list[tuple[int, str]] = []
    for idx, (sline, dline) in enumerate(zip(scan_lines, display_lines), start=1):
        lowered = sline.lower()
        if any(marker in lowered for marker in PEDAGOGICAL_MARKERS):
            continue
        if ADIDAS_PATTERN.search(sline):
            adidas_hits.append((idx, dline.strip()))
    if adidas_hits:
        violations.append(
            f"CAPITALIZED 'Adidas' found on {len(adidas_hits)} line(s) "
            "(adidas is always lowercase, even at sentence start):"
        )
        for ln, ctx in adidas_hits[:5]:
            violations.append(f"    line {ln}: {ctx[:120]}")
        if len(adidas_hits) > 5:
            violations.append(f"    ... and {len(adidas_hits) - 5} more")

    # US market language check: soccer footwear is "cleats" (US market), not
    # "boots" (UK/global). Skip pedagogical lines; blank non-soccer "boot" phrases
    # before matching. Mirrors the Adidas check (scan_lines already have backtick
    # and fenced content blanked).
    boot_hits: list[tuple[int, str]] = []
    for idx, (sline, dline) in enumerate(zip(scan_lines, display_lines), start=1):
        lowered = sline.lower()
        if any(marker in lowered for marker in PEDAGOGICAL_MARKERS):
            continue
        cleaned = BOOT_NON_SOCCER.sub(lambda m: " " * len(m.group(0)), sline)
        if BOOT_PATTERN.search(cleaned):
            boot_hits.append((idx, dline.strip()))
    if boot_hits:
        violations.append(
            f"UK/GLOBAL 'boot(s)' found on {len(boot_hits)} line(s) "
            "(US market term is 'cleat(s)'; 'shoe(s)' acceptable for variation):"
        )
        for ln, ctx in boot_hits[:5]:
            violations.append(f"    line {ln}: {ctx[:120]}")
        if len(boot_hits) > 5:
            violations.append(f"    ... and {len(boot_hits) - 5} more")

    # Editorial body H2 casing check: editorial body H2s use sentence case (first
    # word capitalized; "adidas" the sole lowercase-start exception). Scope-limited
    # to the "### Description" -> "## Product Details" region, so it runs on every
    # file but only ever matches inside a real brief.
    body_h2_hits = find_lowercase_body_h2s(display_lines)
    if body_h2_hits:
        violations.append(
            f"LOWERCASE EDITORIAL BODY H2 found on {len(body_h2_hits)} line(s) "
            "(editorial body H2s use sentence case: capitalize the first word; "
            "'adidas' is the only lowercase-start exception):"
        )
        for ln, ctx in body_h2_hits[:5]:
            violations.append(f"    line {ln}: {ctx[:120]}")
        if len(body_h2_hits) > 5:
            violations.append(f"    ... and {len(body_h2_hits) - 5} more")

    # Internal link format check (deliverables/ and briefings/ only): link
    # suggestions must be full HTTPS canonical-domain URLs. Flag insecure http://
    # ProSoccer URLs and mangled missing-domain links. Skip pedagogical lines;
    # backtick/fenced content is already blanked in scan_lines.
    if link_scope:
        for label, pattern, detail in (
            ("INSECURE LINK", LINK_INSECURE_PATTERN,
             "ProSoccer links must be https:// (insecure http:// found; "
             "use https://www.prosoccer.com/...):"),
            ("MANGLED LINK", LINK_MANGLED_PATTERN,
             "link is missing its domain (e.g. http:///path); "
             "use the full https://www.prosoccer.com/... form:"),
        ):
            link_hits: list[tuple[int, str]] = []
            for idx, (sline, dline) in enumerate(zip(scan_lines, display_lines), start=1):
                lowered = sline.lower()
                if any(marker in lowered for marker in PEDAGOGICAL_MARKERS):
                    continue
                if pattern.search(sline):
                    link_hits.append((idx, dline.strip()))
            if link_hits:
                violations.append(f"{label}: {detail}")
                for ln, ctx in link_hits[:5]:
                    violations.append(f"    line {ln}: {ctx[:120]}")
                if len(link_hits) > 5:
                    violations.append(f"    ... and {len(link_hits) - 5} more")

    return violations


def check(text: str, path: Path | None = None) -> int:
    """Run the voice check against the text. Return 0 if clean, 1 if violations found.

    Thin CLI wrapper over `collect_violations`: same detection, plus the pass/fail
    print output the command-line and existing tests expect.
    """
    violations = collect_violations(text, path)
    if violations:
        print("VOICE CHECK FAILED")
        for line in violations:
            print(f"  {line}")
        return 1

    print("VOICE CHECK PASSED: no em-dashes, no en-dashes, no forbidden words/phrases/openers, no capitalized Adidas, no UK 'boots', no lowercase editorial body H2s.")
    return 0


def main() -> int:
    parser = argparse.ArgumentParser(description="Voice-check a file against the ProSoccer brand-voice forbidden list.")
    parser.add_argument("path", nargs="?", help="Path to the file to check (.md, .txt, or .docx).")
    parser.add_argument("--file", dest="file_flag", help="Alternate way to pass the file path.")
    args = parser.parse_args()

    target = args.path or args.file_flag
    if not target:
        parser.error("A file path is required (positional or --file).")

    path = Path(target).resolve()
    if not path.exists():
        print(f"ERROR: file not found: {path}", file=sys.stderr)
        return 2

    if path.suffix.lower() not in SUPPORTED_EXTS:
        print(f"ERROR: unsupported file type '{path.suffix}'. Supported: {sorted(SUPPORTED_EXTS)}", file=sys.stderr)
        return 2

    text = extract_text(path)
    print(f"Checking: {path}")
    return check(text, path)


if __name__ == "__main__":
    sys.exit(main())
